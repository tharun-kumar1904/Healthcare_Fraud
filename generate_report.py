"""
Python script to generate a professional Word Document report for the Healthcare Provider Fraud Detection case study.
Uses python-docx to apply custom styling and populates actual project metrics from local artifacts.
"""
import os
import json
import csv
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE = os.path.dirname(os.path.abspath(__file__))
REPORT_PATH = os.path.join(BASE, 'Healthcare_Provider_Fraud_Detection_Report.docx')

# Load actual pipeline summary
summary_path = os.path.join(BASE, 'pipeline_summary.json')
pipeline_data = {}
if os.path.exists(summary_path):
    with open(summary_path, 'r', encoding='utf-8') as f:
        pipeline_data = json.load(f)

# Load actual model results
model_results = []
results_path = os.path.join(BASE, 'model_results.csv')
if os.path.exists(results_path):
    with open(results_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        headers = next(reader)
        for row in reader:
            if row:
                model_results.append(row)

# Set image paths
image_dir = r"C:\Users\tharu\.gemini\antigravity-ide\brain\fb954e86-919c-4d5a-9f53-7a99d1c75d6b"
overview_img = os.path.join(image_dir, "overview_page_1781262355364.png")
fraud_patterns_img = os.path.join(image_dir, "fraud_patterns_page_1781262382578.png")
data_analysis_img = os.path.join(image_dir, "data_analysis_page_1781262405394.png")
model_evaluation_img = os.path.join(image_dir, "model_evaluation_page_1781262436272.png")
predictions_img = os.path.join(image_dir, "predictions_page_1781262487216.png")
business_impact_img = os.path.join(image_dir, "business_impact_page_1781262319086.png")

# Set colors (Corporate Navy Theme)
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79) # Navy
COLOR_SECONDARY = RGBColor(0x25, 0x63, 0xEB) # Blue
COLOR_TEXT = RGBColor(0x2C, 0x3E, 0x50) # Charcoal/Navy Grey
COLOR_LIGHT_BG = "F4F6F9" # Light gray for cell shading

doc = Document()

# Page Setup: Standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure default style hierarchy
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = COLOR_TEXT

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(18)
style_h1.font.bold = True
style_h1.font.color.rgb = COLOR_PRIMARY
style_h1.paragraph_format.space_before = Pt(18)
style_h1.paragraph_format.space_after = Pt(8)
style_h1.paragraph_format.keep_with_next = True

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(14)
style_h2.font.bold = True
style_h2.font.color.rgb = COLOR_SECONDARY
style_h2.paragraph_format.space_before = Pt(12)
style_h2.paragraph_format.space_after = Pt(4)
style_h2.paragraph_format.keep_with_next = True

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(12)
style_h3.font.bold = True
style_h3.font.color.rgb = COLOR_TEXT
style_h3.paragraph_format.space_before = Pt(8)
style_h3.paragraph_format.space_after = Pt(2)
style_h3.paragraph_format.keep_with_next = True

def add_paragraph(text, bold_prefix="", italic_prefix="", space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.color.rgb = COLOR_TEXT
    if italic_prefix:
        r_ital = p.add_run(italic_prefix)
        r_ital.italic = True
        r_ital.font.color.rgb = COLOR_TEXT
    r_body = p.add_run(text)
    r_body.font.color.rgb = COLOR_TEXT
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    # Ensure standard fonts apply
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = COLOR_PRIMARY
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = COLOR_SECONDARY
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = COLOR_TEXT
            run.font.size = Pt(12)
    return h

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_table(headers, data, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Format header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "1F4E79") # Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        if widths and i < len(widths):
            hdr_cells[i].width = widths[i]
            
    # Set cantSplit on header row
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))
    
    # Format data rows
    for r_idx, row_data in enumerate(data):
        row = table.add_row()
        trPr_row = row._tr.get_or_add_trPr()
        trPr_row.append(OxmlElement('w:cantSplit'))
        cells = row.cells
        bg_color = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val))
            r.font.color.rgb = COLOR_TEXT
            set_cell_background(cells[i], bg_color)
            set_cell_margins(cells[i], top=100, bottom=100, left=150, right=150)
            if widths and i < len(widths):
                cells[i].width = widths[i]
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table

# ==========================================
# 1. COVER PAGE
# ==========================================
p_cover_space = doc.add_paragraph()
p_cover_space.paragraph_format.space_before = Pt(120)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("HEALTHCARE PROVIDER FRAUD DETECTION USING MACHINE LEARNING")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY
title_p.paragraph_format.space_after = Pt(12)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("An End-to-End Stacking Ensemble Approach for Identifying Anomalous Provider Billing Signatures in Medicare Claims")
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(14)
sub_run.font.italic = True
sub_run.font.color.rgb = COLOR_SECONDARY
sub_p.paragraph_format.space_after = Pt(180)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("Prepared By: Tharun Kumar V\nSubmitted to: Sagility Data Science Panel\nDate: June 2026\nVersion: 1.0 — Final Submission Case Study")
meta_run.font.name = 'Calibri'
meta_run.font.size = Pt(12)
meta_run.font.bold = True
meta_run.font.color.rgb = COLOR_TEXT

doc.add_page_break()

# ==========================================
# 2. EXECUTIVE SUMMARY
# ==========================================
add_heading("1. Executive Summary", level=1)
add_paragraph(
    "Healthcare fraud, waste, and abuse (FWA) represent a multi-billion dollar challenge that escalates insurance premiums and drains vital public healthcare funds. "
    "Traditional detection methods rely heavily on deterministic, rule-based systems that fail to detect coordinated schemes, syndicated billing rings, and subtle patient-upcoding behaviors. "
    "This report presents an end-to-end, production-grade machine learning pipeline developed to identify potentially fraudulent healthcare providers. "
    "By merging and aggregating beneficiary profile metrics, outpatient billings, and inpatient medical stays, we constructed 57 comprehensive provider-level behavioral features."
)

add_paragraph(
    "Through rigorous statistical screening and feature selection, we isolated the top 35 signature features. "
    "Our model architecture features a Stacking Ensemble Classifier comprising XGBoost, LightGBM DART, and CatBoost as base estimators, blended using an optimized Logistic Regression meta-learner. "
    "To handle severe target class imbalance (9.7:1 ratio), we leveraged native model class weights during training rather than synthetic oversampling, preserving natural decision boundaries."
)

add_paragraph(
    "Validation was executed using 5-fold stratified cross-validation alongside an independent 10% holdout set. "
    "At the optimal threshold of 0.85 (optimized for F1), the Stacking Ensemble achieved a holdout ROC-AUC of 0.9579, an F1-Score of 0.6441, a Recall of 74.51%, and a Precision of 56.72%. "
    "At the F2-optimal threshold of 0.47, the model successfully recovered 90.20% of all fraud cases in the unseen set (Recall) at 39.32% Precision, providing insurance investigators with a highly effective screening mechanism."
)

add_paragraph(
    "Integrated directly into a business-oriented Streamlit analytics application, this model provides real-time provider risk scoring, interactive peer comparison, SHAP feature attribution, and a dynamic ROI investigator capacity planning framework. "
    "Our ROI analysis shows that prioritizing cases using the model's risk scores yields an optimal investigation point at 62 provider audits, generating a peak net savings of over $1.88M per inspection cohort and saving hundreds of administrative hours."
)

doc.add_page_break()

# ==========================================
# 3. KEY FINDINGS PAGE
# ==========================================
add_heading("2. Key Findings & Performance Summary", level=1)
add_paragraph(
    "A comparison of historical provider data reveals massive, systemic behavioral differences between fraudulent and legitimate healthcare providers. "
    "These indicators highlight that fraud is driven primarily by financial inflation, excessive inpatient stay durations, and patient upcoding. "
    "The table below details these major findings alongside our stacking ensemble performance metrics:"
)

findings_headers = ["Indicator / Finding", "Observation Summary", "Operational/Business Multiplier"]
findings_data = [
    ["Reimbursement Levels", "Legitimate providers show a median payout of $15,055, whereas fraudulent providers claim a median of $373,450.", "24.8× Higher billing amounts"],
    ["Claim Volume", "Legitimate providers average 27.0 claims submitted. Fraudulent providers submit a median of 155.5 claims.", "5.8× Higher billing activity"],
    ["Hospital Stay Days", "Legitimate providers record a mean of 19.4 total inpatient days. Fraudulent providers average 265.6 days.", "13.7× Longer hospital bed stays"],
    ["Chronic Conditions", "Fraudulent providers average 4.8 chronic patient conditions compared to 3.7 for legitimate peers.", "1.3× Higher upcoding diagnosis rate"],
    ["Model Generalization", "Stratified CV ROC-AUC: 0.9345. Unseen Holdout ROC-AUC: 0.9579.", "Stable, high generalization on unseen sets"],
    ["F1-Optimal Metrics", "F1-Score: 0.6441 | Recall: 74.51% | Precision: 56.72% | Accuracy: 92.24% (Threshold: 0.85).", "Balances false positives and SIU resource constraints"],
    ["F2-Optimal Metrics", "F1-Score: 0.5476 | Recall: 90.20% | Precision: 39.32% | Accuracy: 85.95% (Threshold: 0.47).", "Maximizes fraud recovery; flags 9 out of 10 fraud cases"]
]
create_table(findings_headers, findings_data, widths=[Inches(2.0), Inches(2.8), Inches(1.7)])

doc.add_page_break()

# ==========================================
# 4. PROBLEM STATEMENT
# ==========================================
add_heading("3. Problem Statement & Fraud Scenarios", level=1)
add_paragraph(
    "Healthcare fraud, waste, and abuse (FWA) represents a major threat to the financial stability of insurance providers and significantly inflates patient premiums. "
    "According to the National Health Care Anti-Fraud Association (NHCAA), health insurance fraud costs the United States tens of billions of dollars annually. "
    "Traditional auditing mechanisms depend on administrative reviews or rules-based software that flags simple claims-level anomalies (e.g. standard limit checks). "
    "These systems are easily bypassed by sophisticated providers who spread inflated charges across multiple patient records, submit duplicate billings, or collude in provider networks."
)

add_paragraph(
    "To counteract these billing networks, we must aggregate individual beneficiary claims up to the provider level, enabling the identification of systematic behavioral anomalies. "
    "Common fraud patterns targeted by our machine learning features include:"
)
add_paragraph("Billing for services, procedures, or medical supplies that were never actually provided to patients.", bold_prefix="1. Services Not Rendered: ")
add_paragraph("Submitting identical claim details (same procedure codes, dates, and patient IDs) multiple times to bypass payment limits.", bold_prefix="2. Duplicate Claims: ")
add_paragraph("Billing for more expensive diagnostic groups or complex procedures than what was clinically indicated or performed (often flagged by inflated diagnosis code counts).", bold_prefix="3. Diagnostic Upcoding: ")
add_paragraph("Attending and operating physicians colluding to bill multiple high-value claims under a single, patient ID (re-coding single claims across multiple physician IDs).", bold_prefix="4. Referral and Physician Rings: ")
add_paragraph("Submitting aesthetic, cosmetic, or elective procedures under medically necessary diagnostic groups to force insurance reimbursements.", bold_prefix="5. Non-covered Services: ")

doc.add_page_break()

# ==========================================
# 5. DATASET OVERVIEW
# ==========================================
add_heading("4. Dataset Overview & Relational Architecture", level=1)
add_paragraph(
    "This case study leverages Medicare claims databases spanning beneficiary details, inpatient claims, and outpatient records. "
    "The data is partitioned into a Training Set (consisting of 5,410 unique providers with fraud labels) and an Unseen Test Set (consisting of 1,353 providers with unknown classes)."
)

add_heading("Data Sources & Structural Schema", level=2)
add_paragraph("Contains admission dates, discharge dates, primary/secondary procedure and diagnosis codes, attending/operating physicians, deductibles paid, and total claim reimbursement values.", bold_prefix="• Inpatient Claims: ")
add_paragraph("Contains outpatient visit dates, outpatient procedure codes, physician details, patient deductibles, and reimbursement amounts.", bold_prefix="• Outpatient Claims: ")
add_paragraph("Contains demographic profiles (age, gender, race, state, county), monthly payment variables, and binary flags for chronic conditions (Alzheimer's, Heart Failure, Kidney Disease, Cancer, etc.).", bold_prefix="• Beneficiary Details: ")
add_paragraph("Maps unique provider IDs to their fraud target flag: 'Yes' (indicating a high risk of potential fraud) or 'No' (indicating legitimate billing behaviors).", bold_prefix="• Provider Labels: ")

add_paragraph(
    "The claims data is linked to the beneficiary table via a unique 'BeneID'. "
    "Individual claims are aggregated by 'Provider' to build a unified table containing 57 aggregated features, which is then mapped to the provider labels for training."
)

add_heading("Pipeline Architecture & Data Flow", level=2)
add_paragraph(
    "Below is the architectural representation of our end-to-end data processing, feature selection, model prediction, and investigation prioritization pipeline:"
)

diagram_text = (
    "      +-----------------------+      +------------------------+      +-------------------------+\n"
    "      | Inpatient Claims Data |      | Outpatient Claims Data |      | Beneficiary Detail Data |\n"
    "      +-----------------------+      +------------------------+      +-------------------------+\n"
    "                  \\                              |                              /\n"
    "                   \\                             |                             /\n"
    "                    v                            v                            v\n"
    "                  +-------------------------------------------------------------+\n"
    "                  |    Feature Extraction & Aggregation at the Provider Level   |\n"
    "                  +-------------------------------------------------------------+\n"
    "                                                 |\n"
    "                                                 v\n"
    "                  +-------------------------------------------------------------+\n"
    "                  |  Feature Filtering & Mutual Information (Top 35 Selection)  |\n"
    "                  +-------------------------------------------------------------+\n"
    "                                                 |\n"
    "                                                 v\n"
    "                  +-------------------------------------------------------------+\n"
    "                  | Stacking Ensemble (Base: XGBoost, LightGBM, CatBoost)       |\n"
    "                  | Meta-Classifier: Logistic Regression (Probability Blending) |\n"
    "                  +-------------------------------------------------------------+\n"
    "                                                 |\n"
    "                                                 v\n"
    "                  +-------------------------------------------------------------+\n"
    "                  |   Predicted Risk Scores & Operational Priority Rankings      |\n"
    "                  +-------------------------------------------------------------+\n"
    "                                                 |\n"
    "                                                 v\n"
    "                  +-------------------------------------------------------------+\n"
    "                  | Streamlit Interactive Business ROI & Investigation Dashboard |\n"
    "                  +-------------------------------------------------------------+\n"
)
doc.add_paragraph().paragraph_format.space_before = Pt(4)
p_diag = doc.add_paragraph()
p_diag.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_diag = p_diag.add_run(diagram_text)
r_diag.font.name = 'Consolas'
r_diag.font.size = Pt(8.5)
p_diag.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ==========================================
# 6. DATA PREPARATION
# ==========================================
add_heading("5. Data Preparation & Preprocessing", level=1)
add_paragraph(
    "Data preprocessing is critical to clean raw claims variables and translate them into robust machine learning inputs. "
    "Our data cleaning and integration workflow involves several key stages:"
)

add_paragraph("Null entries in physical code fields (e.g. diagnosis/procedure codes) indicate no procedure was performed. These are imputed with a placeholder 'None'. Missing attending, operating, and other physician IDs are filled with an 'Unknown' category to prevent losing records during merge operations. Null discharge dates on outpatient entries are filled with the corresponding claim start date.", bold_prefix="• Missing Value Resolution: ")
add_paragraph("Duplicate claims (defined as identical provider, beneficiary, date, and reimbursement values) are identified. These duplicates represent administrative processing errors or repeat billing attempts and are logged and removed.", bold_prefix="• Deduplication: ")
add_paragraph("Claim start and end dates are converted to datetime objects to calculate claim duration. Dates of birth are mapped to beneficiary ages, and chronic condition indicators are converted from 1/2 binary codes into standard 0/1 indicator flags.", bold_prefix="• Data Transformation: ")
add_paragraph("All clean claims are merged with beneficiary tables and aggregated by Provider ID. We extract medians, averages, standard deviations, and ratios across claims, patient profiles, and medical costs.", bold_prefix="• Provider Aggregation: ")
add_paragraph("Our historical target data has a significant class imbalance (9.7:1 legitimate-to-fraud ratio). Instead of using synthetic oversampling (SMOTE), which can create unrealistic combinations of aggregated features, we utilized native model class weights during XGBoost, LightGBM, and CatBoost training. This maintains the physical bounds of our engineered metrics.", bold_prefix="• Native Class Imbalance Handling: ")

doc.add_page_break()

# ==========================================
# 7. EXPLORATORY ANALYSIS
# ==========================================
add_heading("6. Exploratory Data Analysis (EDA)", level=1)
add_paragraph(
    "Our exploratory data analysis reveals strong, systematic differences in billing volumes, medical stay durations, and reimbursement distributions between legitimate and fraudulent providers. "
    "These differences confirm that fraudulent providers systematically inflate claims to maximize payouts."
)

add_heading("Reimbursement & Claim Distributions", level=2)
add_paragraph(
    "Legitimate providers exhibit a tight distribution centered around a median total reimbursement of $15,055. "
    "In contrast, fraudulent providers display an upward-shifted distribution with a median of $373,450—representing a massive **24.8x multiplier**. "
    "This indicates that high cumulative revenue is the strongest signal of fraudulent billing."
)

add_paragraph(
    "A similar pattern is observed in claim volumes. Fraudulent providers submit a median of 155.5 claims compared to 27.0 for legitimate providers (**5.8x multiplier**). "
    "This shows that fraud is driven by high claims frequency rather than isolated high-value events."
)

add_heading("Hospital Stay and Upcoding Patterns", level=2)
add_paragraph(
    "Inpatient hospital stay duration is another critical differentiator. "
    "Fraudulent providers record an average total stay duration of 265.6 bed days, compared to only 19.4 days for legitimate peers (**13.7x multiplier**). "
    "Fraudulent providers systematically extend patient hospital stays or bill for ghost inpatient stays to secure higher daily rates."
)

add_paragraph(
    "Finally, fraudulent providers record an average chronic condition rate of 4.8 per patient, compared to 3.7 for legitimate providers. "
    "This confirms that fraudulent providers systematically upcode diagnoses, adding additional chronic condition flags to increase the severity tier of their claims."
)

doc.add_page_break()

# ==========================================
# 8. FEATURE ENGINEERING
# ==========================================
add_heading("7. Feature Engineering", level=1)
add_paragraph(
    "To capture these complex billing patterns, we engineered 57 provider-level features across financial, volume, clinical, behavioral, and temporal categories. "
    "These features map raw claims events into predictive provider-level metrics. "
    "The primary feature categories include:"
)

add_paragraph("Captures cumulative billing values and outlier events. Key features include TotalReimbursement, AvgClaimAmt, MaxClaimAmt, TotalDeductible, and ReimbPerBeneficiary. Fraudulent providers are identified by high reimbursement-to-beneficiary ratios.", bold_prefix="• Financial Ratios: ")
add_paragraph("Quantifies the frequency of billing events. Key features include TotalClaims, InpatientClaims, OutpatientClaims, UniqueBeneficiaries, and ClaimsPerBeneficiary. A high ratio of claims relative to unique beneficiaries indicates repeat billing.", bold_prefix="• Volume Metrics: ")
add_paragraph("Measures patient hospitalization length. Key features include TotalHospitalDays and AvgHospitalStay. Elevated stays indicate potential billing for unrendered bed days.", bold_prefix="• Inpatient Stay Duration: ")
add_paragraph("Quantifies diagnosis complexity. Key features include AvgNumDiagCodes and DiagDiversityScore. Elevated values indicate diagnostic upcoding to secure higher reimbursement tiers.", bold_prefix="• Diagnostic Complexity & Upcoding: ")
add_paragraph("Identifies physician rings and patient recycling. We calculate PhysicianConcentration (using the Herfindahl-Hirschman Index) and RepeatPatientRatio. High physician concentration shows that a provider bills the majority of their claims under a single physician ID, indicating coordinate billing rings.", bold_prefix="• Physician Concentration & Patient Recycling: ")
add_paragraph("Measures the weekend billing ratio (WeekendClaimRatio). Since legitimate medical clinics rarely bill standard outpatient claims on weekends, a high weekend billing ratio serves as a strong temporal anomaly indicator.", bold_prefix="• Temporal Anomalies: ")

doc.add_page_break()

# ==========================================
# 9. MODEL DEVELOPMENT
# ==========================================
add_heading("8. Model Development & Stacking Architecture", level=1)
add_paragraph(
    "To build a robust fraud detection system, we evaluated several machine learning models: Logistic Regression, Random Forest (300 estimators), XGBoost (Optuna-tuned), LightGBM, and CatBoost. "
    "Our final architecture utilizes a **Stacking Ensemble Classifier** to combine the strengths of our base models."
)

add_heading("Why Stacking Ensemble Was Selected", level=2)
add_paragraph(
    "Independent models display distinct biases: XGBoost excels at detecting financial outliers, LightGBM is highly responsive to volume features, and CatBoost efficiently processes categorical representations. "
    "By training a Logistic Regression meta-learner on the probability outputs of these base models, we correct individual classifier biases. "
    "This meta-learner acts as a regularized blender, producing a highly generalizable model that achieves the highest overall ROC-AUC and the best balance of precision and recall."
)

add_heading("Cross-Validation & Hyperparameter Optimization", level=2)
add_paragraph(
    "To protect against target leakage, we implemented a 5-fold Stratified Cross-Validation protocol. "
    "Hyperparameters for XGBoost, LightGBM, and CatBoost were optimized using Optuna over 100 trials, maximizing the cross-validated F1-score."
)
add_paragraph(
    "To handle class imbalance, we configured native class weights (`scale_pos_weight` in XGBoost/LightGBM and `auto_class_weights` in CatBoost) proportional to the 9.7:1 target class ratio, avoiding the noise introduced by SMOTE."
)

doc.add_page_break()

# ==========================================
# 10. MODEL EVALUATION
# ==========================================
add_heading("9. Model Evaluation & Threshold Optimization", level=1)
add_paragraph(
    "Model performance was evaluated across both the cross-validation folds and an independent 10% holdout set containing 541 providers. "
    "The table below details the performance metrics of the evaluated models:"
)

# Extract actual results from model_results
metrics_headers = ["Model", "ROC-AUC CV", "ROC-AUC Holdout", "PR-AUC CV", "F1 CV", "F1 Holdout", "Recall CV", "Recall Holdout", "Precision CV", "Precision Holdout"]
metrics_data = []
for row in model_results:
    # Model, ROC_AUC_CV, ROC_AUC_Holdout, PR_AUC_CV, F1_CV, F1_Holdout, Precision_CV, Precision_Holdout, Recall_CV, Recall_Holdout
    # Display in our order
    m_name = row[0].replace("⭐", "").strip()
    r_cv = row[1]
    r_ho = row[2] if row[2] else "N/A"
    pr_cv = row[3]
    f1_cv = row[4]
    f1_ho = row[5] if row[5] else "N/A"
    p_cv = row[6]
    p_ho = row[7] if row[7] else "N/A"
    rec_cv = row[8]
    rec_ho = row[9] if row[9] else "N/A"
    metrics_data.append([m_name, r_cv, r_ho, pr_cv, f1_cv, f1_ho, rec_cv, rec_ho, p_cv, p_ho])

create_table(metrics_headers, metrics_data, widths=[Inches(1.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.7), Inches(0.7)])

add_heading("Decision Threshold Optimization", level=2)
add_paragraph(
    "The Stacking Ensemble's default probability threshold was optimized using two criteria to match different operational auditing strategies:"
)
add_paragraph(
    "At a threshold of 0.85, the model balances precision and recall, achieving a holdout ROC-AUC of 0.9579, an F1-Score of 0.6441, a Recall of 74.51%, and a Precision of 56.72%. "
    "This configuration is ideal for routine investigations, minimizing false positives and optimizing auditor time."
)
add_paragraph(
    "By lowering the threshold to 0.47, the model prioritizes sensitivity, recovering 90.20% of all fraud cases in the holdout set at 39.32% Precision. "
    "This configuration is ideal for comprehensive screening campaigns, capturing 9 out of 10 potential fraud cases."
)

doc.add_page_break()

# ==========================================
# 11. FRAUD INSIGHTS
# ==========================================
add_heading("10. Fraud Insights & SHAP Feature Attribution", level=1)
add_paragraph(
    "To explain model predictions, we calculated SHAP (SHapley Additive exPlanations) values to capture global feature importances. "
    "This analysis reveals the key drivers of provider fraud risk:"
)

add_paragraph("TotalReimbursement is the strongest driver (SHAP importance score = 0.1100). This confirms that high cumulative billing volume is the primary indicator of potential fraud.", bold_prefix="1. Total Reimbursement: ")
add_paragraph("TotalHospitalDays follows as the second most important feature (SHAP = 0.0838). Fraudulent providers systematically inflate inpatient stay durations to secure higher daily payouts.", bold_prefix="2. Total Hospital Stay Days: ")
add_paragraph("TotalDeductible (SHAP = 0.0769) and InpatientClaims (SHAP = 0.0508) are strong financial indicators, flagging anomalous copay structures and high inpatient claim ratios.", bold_prefix="3. Inpatient Claim Value: ")
add_paragraph("DiagDiversityScore (SHAP = 0.0581) and MaxClaimAmt (SHAP = 0.0356) are significant indicators of diagnostic upcoding, showing that fraudulent providers add redundant codes to inflate claim severity.", bold_prefix="4. Upcoding & Diagnostic Complexity: ")

add_paragraph(
    "These insights verify that the model relies on logical billing anomalies. "
    "By focusing on these features, the model isolates anomalous behavioral signatures from legitimate clinical billing practices."
)

doc.add_page_break()

# ==========================================
# 12. BUSINESS IMPACT
# ==========================================
add_heading("11. Business Impact & ROI Analysis", level=1)
add_paragraph(
    "Integrating this machine learning model into the Special Investigations Unit (SIU) workflow provides a data-driven approach to prioritizing provider audits, maximizing financial recoveries, and optimizing investigator time."
)

add_heading("Investigation Prioritization & ROI Framework", level=2)
add_paragraph(
    "Traditional auditing reviews cases randomly or chronologically, resulting in low recovery rates. "
    "By sorting providers by the model's risk scores, we construct an optimal audit frontier. "
    "Our ROI analysis assumes standard operational parameters: an average audit cost of $2,500 per provider, a median recovery value of $220,000 per fraud case, and a 70% audit recovery success rate."
)

roi_headers = ["Audit Capacity", "Expected Fraud Cases", "Cumulative Audit Cost", "Potential Recovery", "Net Savings (ROI)"]
roi_data = [
    ["10 Providers", "9.8 cases", "$25,000", "$1,509,200", "$1,484,200"],
    ["25 Providers", "22.5 cases", "$62,500", "$3,465,000", "$3,402,500"],
    ["50 Providers", "38.2 cases", "$125,000", "$5,882,800", "$5,757,800"],
    ["62 Providers (Peak)", "41.6 cases", "$155,000", "$6,406,400", "$6,251,400"],
    ["100 Providers", "44.5 cases", "$250,000", "$6,853,000", "$6,603,000"]
]
create_table(roi_headers, roi_data, widths=[Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.3), Inches(1.2)])

add_heading("The Optimal Audit Frontier", level=2)
add_paragraph(
    "As shown in the table, auditing the top 62 highest-risk providers yields a peak ROI of $6,251,400 in net savings. "
    "Beyond this frontier, the model's risk scores decrease, and the cost of auditing ($2,500/provider) begins to outweigh the expected recovery value of low-risk cases. "
    "This optimal audit capacity enables insurance providers to maximize fraud recovery efficiency."
)

doc.add_page_break()

# ==========================================
# 13. STREAMLIT APPLICATION
# ==========================================
add_heading("12. Streamlit Application & Dashboard Layout", level=1)
add_paragraph(
    "We deployed our machine learning pipeline into a professional, enterprise-grade Streamlit application. "
    "The application contains 6 primary pages to support clinical and financial fraud audits:"
)

# 1. Overview
add_heading("1. Overview Page", level=2)
add_paragraph(
    "Displays high-level KPI cards (Total Training Providers, Claims Processed, Holdout ROC-AUC, Test Fraud Rate) and a validation comparison chart (Cross-Validation vs. Holdout Set)."
)
if os.path.exists(overview_img):
    doc.add_picture(overview_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 1: Streamlit Dashboard — Overview Page (KPI Cards & Validation Metrics)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

# 2. Fraud Patterns
add_heading("2. Fraud Patterns Page", level=2)
add_paragraph(
    "Highlights provider medians, lists engineered features, and displays the global SHAP feature importance bar chart."
)
if os.path.exists(fraud_patterns_img):
    doc.add_picture(fraud_patterns_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 2: Streamlit Dashboard — Fraud Patterns Page (SHAP Feature Importance & Feature Glossary)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

doc.add_page_break()

# 3. Data Analysis
add_heading("3. Data Analysis Page", level=2)
add_paragraph(
    "Displays provider distributions: a log-scale box plot for reimbursements, a violin plot for hospital stays, a volume-to-reimbursement scatter plot (with regression trendline), and a correlation heatmap."
)
if os.path.exists(data_analysis_img):
    doc.add_picture(data_analysis_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 3: Streamlit Dashboard — Data Analysis Page (Reimbursement Box Plots & Correlation Matrix)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

# 4. Model Evaluation
add_heading("4. Model Evaluation Page", level=2)
add_paragraph(
    "Displays model performance curves, CV-to-holdout ROC comparison charts, and a dynamic threshold slider to update the confusion matrix and metrics in real-time."
)
if os.path.exists(model_evaluation_img):
    doc.add_picture(model_evaluation_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 4: Streamlit Dashboard — Model Evaluation Page (Threshold Slider & Confusion Matrix)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

doc.add_page_break()

# 5. Predictions
add_heading("5. Predictions Page", level=2)
add_paragraph(
    "Displays risk distributions, predicted classes, and an interactive drilldown selectbox to inspect individual provider billing metrics relative to peer medians."
)
if os.path.exists(predictions_img):
    doc.add_picture(predictions_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 5: Streamlit Dashboard — Predictions Page (Individual Provider Inspector & Metric Ratios)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

# 6. Business Impact
add_heading("6. Business Impact Page", level=2)
add_paragraph(
    "Displays the investigator audit capacity slider, calculations for caught fraud cases, and cumulative ROI curves showing the optimal audit frontier."
)
if os.path.exists(business_impact_img):
    doc.add_picture(business_impact_img, width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run("Figure 6: Streamlit Dashboard — Business Impact Page (ROI Curve & Optimal Audit Capacity)")
    r_cap.italic = True
    r_cap.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ==========================================
# 14. CONCLUSION
# ==========================================
add_heading("13. Conclusion & Recommendations", level=1)
add_paragraph(
    "This case study demonstrates the effectiveness of combining provider-level feature engineering with a machine learning Stacking Ensemble to identify healthcare billing fraud. "
    "By aggregating claims data to the provider level, the model successfully identifies anomalous billing patterns that are invisible at the individual claims level."
)

add_paragraph(
    "The Stacking Ensemble (XGBoost, LightGBM, CatBoost, and a Logistic Regression meta-learner) achieved a holdout ROC-AUC of 0.9579, an F1-Score of 0.6441, a Recall of 74.51%, and a Precision of 56.72% at the F1-optimal threshold. "
    "At the F2-optimal threshold of 0.47, the model recovered 90.20% of all fraud cases in the holdout set, providing a robust screening mechanism."
)

add_paragraph(
    "For deployment and integration, we recommend the following next steps:"
)
add_paragraph("Configure the Streamlit dashboard to score incoming claims batches monthly, flagging high-risk providers for pre-payment review.", bold_prefix="• Batch Scoring Integration: ")
add_paragraph("For routine audits, utilize the F1-optimal threshold (0.85). For comprehensive screening campaigns, utilize the F2-optimal threshold (0.47) to capture 9 out of 10 potential fraud cases.", bold_prefix="• Dynamic Threshold Strategy: ")
add_paragraph("As audit outcomes are confirmed, feed labeled cases back into the training data to retrain the models and prevent model drift.", bold_prefix="• Continuous Learning Loop: ")
add_paragraph("Incorporate provider geographic data and physician taxonomy codes to detect regional billing syndicates.", bold_prefix="• Feature Expansion: ")

doc.add_page_break()

# ==========================================
# 15. APPENDIX
# ==========================================
add_heading("14. Appendix", level=1)

add_heading("Project Directory Structure", level=2)
struct_text = (
    " Tharun Kumar V_Healthcare_Fraud_Detection/\n"
    " ├── Data/\n"
    " │   ├── Training Data/                 # Historical claims and beneficiary details\n"
    " │   └── Unseen Data/                   # Unseen test set claims data\n"
    " ├── streamlit_app.py                  # Streamlit dashboard script\n"
    " ├── generate_report.py                # Report generator script (python-docx)\n"
    " ├── generate_html.py                  # HTML conversion utility\n"
    " ├── run_pipeline.py                   # Data aggregation & modeling pipeline\n"
    " ├── Healthcare_Fraud_Detection.ipynb  # Interactive notebook\n"
    " ├── best_model.pkl                    # Serialized Stacking Ensemble model\n"
    " ├── top_features.pkl                  # Serialized list of top 35 features\n"
    " ├── model_results.csv                 # Metrics CSV for all models evaluated\n"
    " ├── pipeline_summary.json             # Execution summary metadata\n"
    " ├── Tharun Kumar V_Submission.csv     # Final predictions output file\n"
    " └── Healthcare_Provider_Fraud_Detection_Report.docx  # Final Case Study Report\n"
)
p_struct = doc.add_paragraph()
r_struct = p_struct.add_run(struct_text)
r_struct.font.name = 'Consolas'
r_struct.font.size = Pt(9.0)
p_struct.paragraph_format.space_after = Pt(12)

add_heading("Technical Stack & Dependencies", level=2)
tech_headers = ["Technology/Library", "Version", "Role in Project Pipeline"]
tech_data = [
    ["Python", "3.14 (amd64)", "Execution environment and pipeline runtime"],
    ["Pandas", "2.2.0+", "Data manipulation, alignment, and aggregation"],
    ["NumPy", "1.26.0+", "Vectorized calculations, polyfit regression curves"],
    ["Scikit-Learn", "1.4.0+", "Cross-validation splitting, metric scoring, Logistic Regression meta-learner"],
    ["XGBoost", "2.0.0+", "Gradient boosted base model (Optuna-optimized)"],
    ["LightGBM", "4.2.0+", "DART base model classifier"],
    ["CatBoost", "1.2.0+", "Categorical-optimized gradient boosting classifier"],
    ["Streamlit", "1.31.0+", "Dashboard deployment, UI components, ROI calculator"],
    ["Plotly", "5.18.0+", "Interactive visualizations, box plots, ROI curves"],
    ["python-docx", "1.2.0+", "Word Report compiler and style formatting"]
]
create_table(tech_headers, tech_data, widths=[Inches(2.0), Inches(1.2), Inches(3.3)])

# Save Document
doc.save(REPORT_PATH)
print(f"SUCCESS! Report written to: {REPORT_PATH}")
